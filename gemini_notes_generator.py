from google import genai
from google.genai import types
from docx import Document
from datetime import datetime
import os
import re

DOCX_DIR = "generated_docs"
os.makedirs(DOCX_DIR, exist_ok=True)

# 1. ADD 'template: str = ""' TO THE FUNCTION DEFINITION
def generate_structured_notes(transcript_text: str, user_prompt: str = "", custom_title: str = None, template: str = ""):
    """
    Generates structured notes from a transcript using Gemini and exports ONLY to DOCX.
    
    Returns: docx_path, pdf_path (always None), final_title
    """
    if transcript_text.startswith("ERROR:"):
        return None, None, transcript_text 

    api_key = os.getenv("GEMINI_API_KEY")
    client = genai.Client(api_key=api_key)
    print("🧠 Generating structured notes with Gemini...")

    prompt = (
        '''Please carefully review the attached audio file and its transcription. Convert this into clear, organized, and 
        easy-to-understand notes that reflect exactly what was said. Automatically create a logical structure using headings, bullet points, 
        and sections to highlight important parts like key topics, decisions, names, dates, and action items. Add timestamps where useful to 
        connect notes to the audio. Make the notes concise but informative, so they can be quickly read and understood. At the end, provide 
        a brief summary that captures the most important points and takeaways from the recording. Format the notes in a way that is neat and 
        professional, adapting the structure naturally based on the content of the audio and transcription. Ensure the notes serve as a reliable, 
        user-friendly document that anyone can easily follow and use.\n\n'''
        f"Transcript:\n{transcript_text}"
    )

    # 2. ADD LOGIC TO USE THE 'template' VARIABLE
    if template:
        prompt += f"\n\nUse this note style template: {template}."
    elif user_prompt:
         prompt += f"\n\nUser's prompt: {user_prompt}"


    try:
        # Use a model that is in your list: "models/gemini-1.5-flash"
        # Note: Your provided file had "gemini-2.5-flash", which may not exist.
        # "gemini-1.5-flash" is safer.
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
        )
        notes_text = response.candidates[0].content.parts[0].text.strip()
    except Exception as e:
        print(f"Error generating structured notes (Gemini call): {e}")
        # Try to provide a more specific error for model not found
        if "model" in str(e) and "not found" in str(e):
             return None, None, f"ERROR: The AI Model specified is incorrect or unavailable: {e}"
        return None, None, f"ERROR: Note Generation API call failed: {e}"

    if custom_title:
        sanitized_title = re.sub(r'[^\w\-]', '', custom_title.replace(' ', '-'))
    else:
        sanitized_title = "AI_Notes_" + datetime.now().strftime("%Y%m%d_%H%M%S")
        
    final_title = sanitized_title

    docx_path = os.path.join(DOCX_DIR, f"{final_title}.docx")
    pdf_path = None

    try:
        document = Document()
        document.add_heading(final_title.replace('_', ' ').replace('-', ' '), level=1)
        for line in notes_text.split("\n"):
            stripped_line = line.strip()
            if stripped_line.startswith(("•", "-")):
                document.add_paragraph(stripped_line, style="List Bullet")
            elif stripped_line:
                document.add_paragraph(stripped_line)
        document.save(docx_path)
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return None, None, f"ERROR: Failed to save document to DOCX: {e}"
        
    print(f"✅ Notes generated (DOCX only): {final_title}")
    return docx_path, pdf_path, final_title